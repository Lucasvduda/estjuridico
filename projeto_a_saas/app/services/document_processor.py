"""
LegalShield AI 2026 — Document Processor
Extração de texto de PDF (com OCR), DOCX e TXT com sanitização obrigatória.
"""

import io
import os
import re
import hashlib
import logging
from enum import Enum
from pathlib import Path
from typing import Optional

import chardet
import fitz  # PyMuPDF
from docx import Document as DocxDocument
from PIL import Image
from pydantic import BaseModel, Field

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Constantes de segurança
# ---------------------------------------------------------------------------
MAX_FILE_SIZE_BYTES = 50 * 1024 * 1024  # 50 MB
ALLOWED_MIME_TYPES = {
    "application/pdf": "pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
    "text/plain": "txt",
}
ALLOWED_EXTENSIONS = {".pdf", ".docx", ".txt"}


class FileType(str, Enum):
    PDF = "pdf"
    DOCX = "docx"
    TXT = "txt"


class ProcessedDocument(BaseModel):
    """Resultado do processamento de um documento."""
    filename: str
    file_type: FileType
    text: str
    page_count: int = 1
    ocr_used: bool = False
    sha256_hash: str
    char_count: int = Field(default=0)
    word_count: int = Field(default=0)

    def model_post_init(self, __context) -> None:
        self.char_count = len(self.text)
        self.word_count = len(self.text.split())


class DocumentProcessingError(Exception):
    """Erro durante processamento de documento."""
    pass


class FileSecurityError(Exception):
    """Erro de segurança na validação do arquivo."""
    pass


# ---------------------------------------------------------------------------
# Sanitização e Validação
# ---------------------------------------------------------------------------

def _detect_mime_type(file_bytes: bytes) -> str:
    """Detecta MIME type real do arquivo baseado em magic bytes."""
    # PDF: começa com %PDF
    if file_bytes[:5] == b"%PDF-":
        return "application/pdf"
    # DOCX: é um ZIP com content_types
    if file_bytes[:4] == b"PK\x03\x04":
        # Verificar se contém [Content_Types].xml (característica de DOCX)
        if b"word/" in file_bytes[:2000] or b"[Content_Types].xml" in file_bytes[:2000]:
            return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        return "application/zip"
    # TXT: verificar se é texto válido
    try:
        file_bytes[:1000].decode("utf-8")
        return "text/plain"
    except UnicodeDecodeError:
        try:
            detected = chardet.detect(file_bytes[:1000])
            if detected["confidence"] and detected["confidence"] > 0.5:
                return "text/plain"
        except Exception:
            pass
    return "application/octet-stream"


def _sanitize_filename(filename: str) -> str:
    """Remove caracteres perigosos do nome do arquivo."""
    # Remove path traversal
    filename = os.path.basename(filename)
    # Mantém apenas caracteres seguros
    filename = re.sub(r'[^\w\s\-.]', '_', filename)
    # Remove espaços múltiplos
    filename = re.sub(r'\s+', '_', filename)
    return filename[:255]  # Limita tamanho


def validate_file(
    file_bytes: bytes,
    filename: str,
) -> tuple[FileType, str]:
    """
    Valida arquivo antes do processamento.
    Retorna (FileType, filename_sanitizado) ou levanta exceção.
    """
    # 1. Verificar tamanho
    if len(file_bytes) > MAX_FILE_SIZE_BYTES:
        raise FileSecurityError(
            f"Arquivo excede limite de {MAX_FILE_SIZE_BYTES // (1024*1024)}MB"
        )

    if len(file_bytes) == 0:
        raise FileSecurityError("Arquivo vazio")

    # 2. Verificar extensão
    safe_filename = _sanitize_filename(filename)
    ext = Path(safe_filename).suffix.lower()
    if ext not in ALLOWED_EXTENSIONS:
        raise FileSecurityError(
            f"Extensão '{ext}' não permitida. Use: {', '.join(ALLOWED_EXTENSIONS)}"
        )

    # 3. Verificar MIME type real (magic bytes)
    detected_mime = _detect_mime_type(file_bytes)
    if detected_mime not in ALLOWED_MIME_TYPES:
        raise FileSecurityError(
            f"Tipo de arquivo detectado '{detected_mime}' não é permitido. "
            f"Extensão diz '{ext}' mas conteúdo real é diferente."
        )

    # 4. Verificar consistência extensão x MIME
    expected_ext = f".{ALLOWED_MIME_TYPES[detected_mime]}"
    if ext != expected_ext:
        raise FileSecurityError(
            f"Inconsistência: extensão '{ext}' não corresponde ao conteúdo real "
            f"(detectado como {detected_mime})"
        )

    file_type = FileType(ALLOWED_MIME_TYPES[detected_mime])
    return file_type, safe_filename


def _compute_hash(file_bytes: bytes) -> str:
    """Calcula SHA-256 do arquivo."""
    return hashlib.sha256(file_bytes).hexdigest()


# ---------------------------------------------------------------------------
# Extração de Texto
# ---------------------------------------------------------------------------

def _extract_pdf_text(file_bytes: bytes) -> tuple[str, int, bool]:
    """
    Extrai texto de PDF. Usa OCR se página não tiver texto selecionável.
    Retorna (texto, num_paginas, usou_ocr).
    """
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    all_text = []
    ocr_used = False
    page_count = len(doc)

    for page_num in range(page_count):
        page = doc[page_num]
        text = page.get_text("text").strip()

        # Se página tem pouco texto, tentar OCR
        if len(text) < 50:
            try:
                ocr_text = _ocr_page(page)
                if len(ocr_text) > len(text):
                    text = ocr_text
                    ocr_used = True
            except Exception as e:
                logger.warning(
                    "OCR falhou na página %d: %s", page_num + 1, str(e)
                )

        if text:
            all_text.append(f"--- Página {page_num + 1} ---\n{text}")

    doc.close()
    return "\n\n".join(all_text), page_count, ocr_used


def _ocr_page(page: fitz.Page) -> str:
    """Aplica OCR em uma página do PDF usando pytesseract."""
    try:
        import pytesseract
    except ImportError:
        logger.warning("pytesseract não instalado. OCR indisponível.")
        return ""

    # Renderiza página como imagem em alta resolução
    mat = fitz.Matrix(2.0, 2.0)  # 2x zoom para melhor OCR
    pix = page.get_pixmap(matrix=mat)
    img_bytes = pix.tobytes("png")
    image = Image.open(io.BytesIO(img_bytes))

    # OCR com Tesseract (português + inglês)
    text = pytesseract.image_to_string(image, lang="por+eng")
    return text.strip()


def _extract_docx_text(file_bytes: bytes) -> str:
    """Extrai texto de arquivo DOCX."""
    doc = DocxDocument(io.BytesIO(file_bytes))
    paragraphs = []

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            paragraphs.append(text)

    # Extrair também texto de tabelas
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(
                cell.text.strip() for cell in row.cells if cell.text.strip()
            )
            if row_text:
                paragraphs.append(f"[Tabela] {row_text}")

    return "\n\n".join(paragraphs)


def _extract_txt_text(file_bytes: bytes) -> str:
    """Extrai texto de arquivo TXT com detecção de encoding."""
    detected = chardet.detect(file_bytes)
    encoding = detected.get("encoding", "utf-8") or "utf-8"

    try:
        return file_bytes.decode(encoding).strip()
    except (UnicodeDecodeError, LookupError):
        # Fallback para latin-1 (aceita qualquer byte)
        return file_bytes.decode("latin-1").strip()


# ---------------------------------------------------------------------------
# API Pública
# ---------------------------------------------------------------------------

def process_document(
    file_bytes: bytes,
    filename: str,
) -> ProcessedDocument:
    """
    Processa um documento completo: valida, extrai texto, retorna resultado estruturado.

    Args:
        file_bytes: Conteúdo bruto do arquivo
        filename: Nome original do arquivo

    Returns:
        ProcessedDocument com texto extraído e metadados

    Raises:
        FileSecurityError: Se arquivo não passar na validação
        DocumentProcessingError: Se extração de texto falhar
    """
    # 1. Validar arquivo
    file_type, safe_filename = validate_file(file_bytes, filename)
    file_hash = _compute_hash(file_bytes)

    logger.info(
        "Processando documento",
        extra={
            "filename": safe_filename,
            "file_type": file_type.value,
            "size_bytes": len(file_bytes),
            "sha256": file_hash[:16] + "...",
        },
    )

    # 2. Extrair texto conforme tipo
    try:
        ocr_used = False
        page_count = 1

        if file_type == FileType.PDF:
            text, page_count, ocr_used = _extract_pdf_text(file_bytes)
        elif file_type == FileType.DOCX:
            text = _extract_docx_text(file_bytes)
        elif file_type == FileType.TXT:
            text = _extract_txt_text(file_bytes)
        else:
            raise DocumentProcessingError(f"Tipo não suportado: {file_type}")

    except (FileSecurityError, DocumentProcessingError):
        raise
    except Exception as e:
        raise DocumentProcessingError(
            f"Erro ao extrair texto de '{safe_filename}': {str(e)}"
        ) from e

    if not text or len(text.strip()) < 10:
        raise DocumentProcessingError(
            f"Nenhum texto extraído de '{safe_filename}'. "
            "Verifique se o arquivo não está corrompido ou em branco."
        )

    return ProcessedDocument(
        filename=safe_filename,
        file_type=file_type,
        text=text,
        page_count=page_count,
        ocr_used=ocr_used,
        sha256_hash=file_hash,
    )
